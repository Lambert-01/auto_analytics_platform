"""
Automated Report Generation Engine for Analytics Platform
Implements: HTML, PDF, Word Report Generation with Templates and Visualizations
"""

import numpy as np
import pandas as pd
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple, Union
import json
import base64
import io
from pathlib import Path
import logging
from dataclasses import dataclass
import warnings
warnings.filterwarnings('ignore')

# Report generation
from jinja2 import Environment, FileSystemLoader, Template
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import plotly.io as pio

# PDF generation
try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Local imports
from app.utils.logger import get_logger
from app.config import settings

logger = get_logger(__name__)


@dataclass
class ReportConfig:
    """Configuration for report generation."""
    # Output formats
    generate_html: bool = True
    generate_pdf: bool = True
    generate_word: bool = True
    
    # Template settings
    template_dir: str = "templates/reports"
    custom_css: Optional[str] = None
    
    # Visualization settings
    figure_width: int = 10
    figure_height: int = 6
    figure_dpi: int = 300
    color_palette: str = "viridis"
    
    # Content settings
    include_executive_summary: bool = True
    include_data_overview: bool = True
    include_visualizations: bool = True
    include_statistical_analysis: bool = True
    include_insights: bool = True
    include_recommendations: bool = True
    include_appendix: bool = True
    
    # Branding
    company_name: str = "NISR Rwanda Analytics Platform"
    logo_path: Optional[str] = None
    primary_color: str = "#1f77b4"
    secondary_color: str = "#ff7f0e"
    
    # Language and formatting
    language: str = "en"
    date_format: str = "%B %d, %Y"
    number_format: str = ",.2f"


class AutomatedReportGenerator:
    """Comprehensive automated report generation system."""
    
    def __init__(self, config: Optional[ReportConfig] = None):
        """Initialize the report generator."""
        self.config = config or ReportConfig()
        self.logger = get_logger(__name__)
        
        # Setup template environment
        self._setup_templates()
        
        # Setup matplotlib and seaborn styles
        self._setup_visualization_styles()
        
        # Report metadata
        self.current_report_data = {}
        
    def _setup_templates(self):
        """Setup Jinja2 template environment."""
        try:
            template_path = Path(self.config.template_dir)
            template_path.mkdir(parents=True, exist_ok=True)
            
            self.template_env = Environment(
                loader=FileSystemLoader(str(template_path)),
                autoescape=True
            )
            
            # Create default templates if they don't exist
            self._create_default_templates()
            
        except Exception as e:
            self.logger.warning(f"Template setup failed: {e}")
            self.template_env = Environment()
    
    def _setup_visualization_styles(self):
        """Setup visualization styles for consistent report appearance."""
        # Matplotlib style
        plt.style.use('seaborn-v0_8')
        plt.rcParams.update({
            'figure.figsize': (self.config.figure_width, self.config.figure_height),
            'figure.dpi': self.config.figure_dpi,
            'font.size': 10,
            'axes.titlesize': 12,
            'axes.labelsize': 10,
            'xtick.labelsize': 9,
            'ytick.labelsize': 9,
            'legend.fontsize': 9,
            'figure.titlesize': 14
        })
        
        # Set color palette
        sns.set_palette(self.config.color_palette)
        
        # Plotly theme
        pio.templates.default = "plotly_white"
    
    def generate_comprehensive_report(self,
                                    dataset_info: Dict[str, Any],
                                    profiling_results: Dict[str, Any],
                                    analysis_results: Dict[str, Any],
                                    ai_insights: Dict[str, Any],
                                    model_results: Optional[Dict[str, Any]] = None,
                                    report_title: str = "Data Analysis Report",
                                    output_dir: str = "./reports") -> Dict[str, str]:
        """
        Generate comprehensive analytical report in multiple formats.
        
        Args:
            dataset_info: Basic dataset information
            profiling_results: Data profiling results
            analysis_results: EDA and statistical analysis results
            ai_insights: AI-generated insights and narratives
            model_results: Optional ML model results
            report_title: Title for the report
            output_dir: Directory to save reports
            
        Returns:
            Dictionary with paths to generated reports
        """
        self.logger.info(f"Generating comprehensive report: {report_title}")
        
        try:
            # Prepare report data
            report_data = self._prepare_report_data(
                dataset_info, profiling_results, analysis_results, 
                ai_insights, model_results, report_title
            )
            
            # Store for reference
            self.current_report_data = report_data
            
            # Generate visualizations
            visualizations = self._generate_report_visualizations(
                dataset_info, profiling_results, analysis_results
            )
            report_data['visualizations'] = visualizations
            
            # Create output directory
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            # Generate reports in different formats
            generated_reports = {}
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"{report_title.replace(' ', '_')}_{timestamp}"
            
            # HTML Report
            if self.config.generate_html:
                html_path = output_path / f"{base_filename}.html"
                self._generate_html_report(report_data, html_path)
                generated_reports['html'] = str(html_path)
            
            # PDF Report
            if self.config.generate_pdf and (WEASYPRINT_AVAILABLE or REPORTLAB_AVAILABLE):
                pdf_path = output_path / f"{base_filename}.pdf"
                self._generate_pdf_report(report_data, pdf_path)
                generated_reports['pdf'] = str(pdf_path)
            
            # Word Report
            if self.config.generate_word and DOCX_AVAILABLE:
                word_path = output_path / f"{base_filename}.docx"
                self._generate_word_report(report_data, word_path)
                generated_reports['word'] = str(word_path)
            
            # Generate report metadata
            metadata = self._generate_report_metadata(report_data, generated_reports)
            metadata_path = output_path / f"{base_filename}_metadata.json"
            with open(metadata_path, 'w') as f:
                json.dump(metadata, f, indent=2)
            
            self.logger.info(f"Report generation completed. Generated {len(generated_reports)} formats.")
            return generated_reports
            
        except Exception as e:
            self.logger.error(f"Error generating report: {e}")
            raise
    
    def _prepare_report_data(self, dataset_info: Dict, profiling_results: Dict,
                           analysis_results: Dict, ai_insights: Dict,
                           model_results: Optional[Dict], report_title: str) -> Dict[str, Any]:
        """Prepare and structure data for report generation."""
        
        report_data = {
            'metadata': {
                'title': report_title,
                'generated_at': datetime.now(),
                'generated_by': self.config.company_name,
                'version': '1.0.0'
            },
            'dataset_info': dataset_info,
            'executive_summary': self._create_executive_summary(
                dataset_info, profiling_results, ai_insights
            ),
            'data_overview': self._create_data_overview(dataset_info, profiling_results),
            'profiling_results': profiling_results,
            'analysis_results': analysis_results,
            'ai_insights': ai_insights,
            'statistical_summary': self._create_statistical_summary(profiling_results),
            'key_findings': self._extract_key_findings(ai_insights),
            'recommendations': self._extract_recommendations(ai_insights),
            'quality_assessment': self._create_quality_assessment(profiling_results),
            'correlation_analysis': self._create_correlation_summary(analysis_results),
            'distribution_analysis': self._create_distribution_summary(profiling_results),
            'model_results': model_results,
            'appendix': self._create_appendix(dataset_info, profiling_results)
        }
        
        return report_data
    
    def _create_executive_summary(self, dataset_info: Dict, profiling_results: Dict, ai_insights: Dict) -> Dict[str, Any]:
        """Create executive summary section."""
        shape = dataset_info.get('shape', (0, 0))
        missing_pct = profiling_results.get('overview', {}).get('missing_data', {}).get('missing_percentage', 0)
        
        # Get AI-generated summary or create fallback
        ai_summary = ai_insights.get('executive_summary', '')
        if not ai_summary:
            ai_summary = f"Analysis of dataset containing {shape[0]:,} records and {shape[1]} features reveals moderate data quality with {missing_pct:.1f}% missing values."
        
        return {
            'ai_generated_summary': ai_summary,
            'key_metrics': {
                'total_records': f"{shape[0]:,}",
                'total_features': shape[1],
                'data_quality_score': self._calculate_quality_score(profiling_results),
                'missing_data_percentage': f"{missing_pct:.1f}%",
                'analysis_timestamp': datetime.now().strftime(self.config.date_format)
            },
            'highlights': ai_insights.get('key_findings', [])[:3]  # Top 3 findings
        }
    
    def _create_data_overview(self, dataset_info: Dict, profiling_results: Dict) -> Dict[str, Any]:
        """Create data overview section."""
        shape = dataset_info.get('shape', (0, 0))
        overview = profiling_results.get('overview', {})
        
        return {
            'basic_info': {
                'rows': f"{shape[0]:,}",
                'columns': shape[1],
                'memory_usage': f"{overview.get('memory_usage', {}).get('total_mb', 0):.1f} MB",
                'file_size_estimate': f"{overview.get('memory_usage', {}).get('total_mb', 0) * 0.8:.1f} MB"
            },
            'column_types': overview.get('column_types', {}),
            'missing_data_summary': overview.get('missing_data', {}),
            'duplicate_data_summary': overview.get('duplicates', {}),
            'data_completeness': self._calculate_completeness_metrics(profiling_results)
        }
    
    def _create_statistical_summary(self, profiling_results: Dict) -> Dict[str, Any]:
        """Create statistical summary section."""
        columns = profiling_results.get('columns', {})
        
        numeric_summaries = {}
        categorical_summaries = {}
        
        for col_name, col_data in columns.items():
            if 'numeric_stats' in col_data:
                numeric_summaries[col_name] = {
                    'descriptive': col_data['numeric_stats'].get('descriptive', {}),
                    'distribution': col_data['numeric_stats'].get('shape', {}),
                    'quartiles': col_data['numeric_stats'].get('quantiles', {})
                }
            
            elif 'categorical_stats' in col_data:
                categorical_summaries[col_name] = {
                    'cardinality': col_data['categorical_stats'].get('cardinality', 0),
                    'entropy': col_data['categorical_stats'].get('entropy', 0),
                    'top_values': col_data['categorical_stats'].get('top_values', {}),
                    'distribution_info': col_data['categorical_stats'].get('value_distribution', {})
                }
        
        return {
            'numeric_columns': numeric_summaries,
            'categorical_columns': categorical_summaries,
            'summary_statistics': self._create_summary_table(numeric_summaries)
        }
    
    def _create_quality_assessment(self, profiling_results: Dict) -> Dict[str, Any]:
        """Create data quality assessment section."""
        overview = profiling_results.get('overview', {})
        columns = profiling_results.get('columns', {})
        
        # Calculate quality metrics
        total_missing = overview.get('missing_data', {}).get('missing_percentage', 0)
        total_duplicates = overview.get('duplicates', {}).get('percentage', 0)
        
        # Column-level quality issues
        quality_issues = []
        high_quality_columns = []
        
        for col_name, col_data in columns.items():
            basic_stats = col_data.get('basic_stats', {})
            missing_pct = basic_stats.get('missing_percentage', 0)
            unique_pct = basic_stats.get('unique_percentage', 0)
            
            if missing_pct > 50:
                quality_issues.append(f"{col_name}: {missing_pct:.1f}% missing values")
            elif missing_pct == 0 and unique_pct > 90:
                high_quality_columns.append(col_name)
        
        return {
            'overall_score': self._calculate_quality_score(profiling_results),
            'missing_data_score': max(0, 100 - total_missing * 2),
            'duplicate_data_score': max(0, 100 - total_duplicates * 5),
            'completeness_score': 100 - total_missing,
            'quality_issues': quality_issues,
            'high_quality_columns': high_quality_columns,
            'recommendations': self._generate_quality_recommendations(quality_issues)
        }
    
    def _create_correlation_summary(self, analysis_results: Dict) -> Dict[str, Any]:
        """Create correlation analysis summary."""
        relationships = analysis_results.get('relationships', {})
        correlations = relationships.get('correlations', {})
        
        return {
            'strong_correlations': correlations.get('strong_correlations', []),
            'correlation_clusters': correlations.get('correlation_clusters', {}),
            'summary_stats': {
                'total_correlations': len(correlations.get('strong_correlations', [])),
                'highest_correlation': self._get_highest_correlation(correlations),
                'multicollinearity_risk': self._assess_multicollinearity_risk(correlations)
            }
        }
    
    def _create_distribution_summary(self, profiling_results: Dict) -> Dict[str, Any]:
        """Create distribution analysis summary."""
        columns = profiling_results.get('columns', {})
        
        distribution_summary = {
            'normal_distributions': [],
            'skewed_distributions': [],
            'high_entropy_categoricals': [],
            'imbalanced_categoricals': []
        }
        
        for col_name, col_data in columns.items():
            if 'numeric_stats' in col_data:
                shape_info = col_data['numeric_stats'].get('shape', {})
                is_normal = shape_info.get('is_normal', False)
                skewness = shape_info.get('skewness', 0)
                
                if is_normal:
                    distribution_summary['normal_distributions'].append(col_name)
                elif abs(skewness) > 2:
                    direction = "right" if skewness > 0 else "left"
                    distribution_summary['skewed_distributions'].append(f"{col_name} ({direction})")
            
            elif 'categorical_stats' in col_data:
                entropy = col_data['categorical_stats'].get('entropy', 0)
                concentration = col_data['categorical_stats'].get('value_distribution', {}).get('most_frequent_percentage', 0)
                
                if entropy > 3:
                    distribution_summary['high_entropy_categoricals'].append(col_name)
                elif concentration > 80:
                    distribution_summary['imbalanced_categoricals'].append(col_name)
        
        return distribution_summary
    
    def _generate_report_visualizations(self, dataset_info: Dict, profiling_results: Dict, analysis_results: Dict) -> Dict[str, str]:
        """Generate visualizations for the report."""
        visualizations = {}
        
        try:
            # Data overview visualization
            visualizations['data_overview'] = self._create_data_overview_chart(profiling_results)
            
            # Missing data visualization
            visualizations['missing_data'] = self._create_missing_data_chart(profiling_results)
            
            # Data types distribution
            visualizations['data_types'] = self._create_data_types_chart(profiling_results)
            
            # Correlation heatmap (if correlations exist)
            correlations = analysis_results.get('relationships', {}).get('correlations', {})
            if correlations.get('pearson_matrix'):
                visualizations['correlation_heatmap'] = self._create_correlation_heatmap(correlations)
            
            # Distribution plots for numeric columns
            visualizations['distributions'] = self._create_distribution_plots(profiling_results)
            
            # Quality score visualization
            visualizations['quality_scores'] = self._create_quality_dashboard(profiling_results)
            
        except Exception as e:
            self.logger.warning(f"Error generating visualizations: {e}")
        
        return visualizations
    
    def _create_data_overview_chart(self, profiling_results: Dict) -> str:
        """Create data overview chart."""
        overview = profiling_results.get('overview', {})
        column_types = overview.get('column_types', {})
        
        if not column_types:
            return ""
        
        # Create pie chart of column types
        fig = go.Figure(data=[go.Pie(
            labels=list(column_types.keys()),
            values=list(column_types.values()),
            hole=0.3,
            marker_colors=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd']
        )])
        
        fig.update_layout(
            title="Distribution of Column Types",
            height=400,
            showlegend=True
        )
        
        return self._plotly_to_base64(fig)
    
    def _create_missing_data_chart(self, profiling_results: Dict) -> str:
        """Create missing data visualization."""
        columns = profiling_results.get('columns', {})
        
        if not columns:
            return ""
        
        # Extract missing data percentages
        missing_data = []
        column_names = []
        
        for col_name, col_data in columns.items():
            basic_stats = col_data.get('basic_stats', {})
            missing_pct = basic_stats.get('missing_percentage', 0)
            
            if missing_pct > 0:  # Only include columns with missing data
                missing_data.append(missing_pct)
                column_names.append(col_name)
        
        if not missing_data:
            # No missing data
            fig = go.Figure()
            fig.add_annotation(
                text="No Missing Data Detected",
                xref="paper", yref="paper",
                x=0.5, y=0.5, xanchor='center', yanchor='middle',
                showarrow=False,
                font=dict(size=20, color="green")
            )
            fig.update_layout(
                title="Missing Data Analysis",
                height=300,
                xaxis={'visible': False},
                yaxis={'visible': False}
            )
        else:
            # Create bar chart of missing data
            fig = go.Figure(data=[go.Bar(
                x=column_names,
                y=missing_data,
                marker_color='red',
                opacity=0.7
            )])
            
            fig.update_layout(
                title="Missing Data by Column",
                xaxis_title="Columns",
                yaxis_title="Missing Percentage (%)",
                height=400
            )
        
        return self._plotly_to_base64(fig)
    
    def _create_data_types_chart(self, profiling_results: Dict) -> str:
        """Create data types distribution chart."""
        overview = profiling_results.get('overview', {})
        column_types = overview.get('column_types', {})
        
        if not column_types:
            return ""
        
        # Create horizontal bar chart
        fig = go.Figure(data=[go.Bar(
            x=list(column_types.values()),
            y=list(column_types.keys()),
            orientation='h',
            marker_color=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728'][:len(column_types)]
        )])
        
        fig.update_layout(
            title="Column Types Distribution",
            xaxis_title="Number of Columns",
            yaxis_title="Data Type",
            height=300
        )
        
        return self._plotly_to_base64(fig)
    
    def _create_correlation_heatmap(self, correlations: Dict) -> str:
        """Create correlation heatmap."""
        pearson_matrix = correlations.get('pearson_matrix', {})
        
        if not pearson_matrix:
            return ""
        
        # Convert to DataFrame for plotting
        df_corr = pd.DataFrame(pearson_matrix)
        
        # Create heatmap
        fig = go.Figure(data=go.Heatmap(
            z=df_corr.values,
            x=df_corr.columns,
            y=df_corr.index,
            colorscale='RdBu',
            zmid=0,
            colorbar=dict(title="Correlation")
        ))
        
        fig.update_layout(
            title="Feature Correlation Heatmap",
            height=500,
            width=500
        )
        
        return self._plotly_to_base64(fig)
    
    def _create_distribution_plots(self, profiling_results: Dict) -> str:
        """Create distribution plots for numeric columns."""
        columns = profiling_results.get('columns', {})
        
        # Find numeric columns
        numeric_columns = []
        for col_name, col_data in columns.items():
            if 'numeric_stats' in col_data:
                numeric_columns.append((col_name, col_data['numeric_stats']))
        
        if not numeric_columns:
            return ""
        
        # Create subplots for distributions
        n_cols = min(2, len(numeric_columns))
        n_rows = (len(numeric_columns) + 1) // 2
        
        fig = make_subplots(
            rows=n_rows, cols=n_cols,
            subplot_titles=[col[0] for col in numeric_columns[:4]],  # Limit to 4 plots
            vertical_spacing=0.1
        )
        
        for i, (col_name, stats) in enumerate(numeric_columns[:4]):
            row = (i // n_cols) + 1
            col = (i % n_cols) + 1
            
            # Create histogram data (simulated since we don't have raw data)
            descriptive = stats.get('descriptive', {})
            mean_val = descriptive.get('mean', 0)
            std_val = descriptive.get('std', 1)
            
            # Generate sample data for visualization
            sample_data = np.random.normal(mean_val, std_val, 1000)
            
            fig.add_trace(
                go.Histogram(x=sample_data, name=col_name, showlegend=False),
                row=row, col=col
            )
        
        fig.update_layout(
            title="Distribution Plots (Sample Data)",
            height=400 * n_rows,
            showlegend=False
        )
        
        return self._plotly_to_base64(fig)
    
    def _create_quality_dashboard(self, profiling_results: Dict) -> str:
        """Create data quality dashboard."""
        quality_score = self._calculate_quality_score(profiling_results)
        overview = profiling_results.get('overview', {})
        
        missing_pct = overview.get('missing_data', {}).get('missing_percentage', 0)
        duplicate_pct = overview.get('duplicates', {}).get('percentage', 0)
        completeness = 100 - missing_pct
        
        # Create gauge charts
        fig = make_subplots(
            rows=1, cols=3,
            specs=[[{"type": "indicator"}, {"type": "indicator"}, {"type": "indicator"}]],
            subplot_titles=["Overall Quality", "Data Completeness", "Duplicate Rate"]
        )
        
        # Overall quality gauge
        fig.add_trace(go.Indicator(
            mode="gauge+number",
            value=quality_score,
            domain={'x': [0, 1], 'y': [0, 1]},
            title={'text': "Quality Score"},
            gauge={
                'axis': {'range': [None, 100]},
                'bar': {'color': "darkblue"},
                'steps': [
                    {'range': [0, 50], 'color': "lightgray"},
                    {'range': [50, 80], 'color': "gray"}
                ],
                'threshold': {
                    'line': {'color': "red", 'width': 4},
                    'thickness': 0.75,
                    'value': 90
                }
            }
        ), row=1, col=1)
        
        # Completeness gauge
        fig.add_trace(go.Indicator(
            mode="gauge+number",
            value=completeness,
            domain={'x': [0, 1], 'y': [0, 1]},
            title={'text': "Completeness %"},
            gauge={
                'axis': {'range': [None, 100]},
                'bar': {'color': "green"},
            }
        ), row=1, col=2)
        
        # Duplicate rate gauge
        fig.add_trace(go.Indicator(
            mode="gauge+number",
            value=duplicate_pct,
            domain={'x': [0, 1], 'y': [0, 1]},
            title={'text': "Duplicate %"},
            gauge={
                'axis': {'range': [None, 20]},
                'bar': {'color': "red"},
            }
        ), row=1, col=3)
        
        fig.update_layout(
            title="Data Quality Dashboard",
            height=300
        )
        
        return self._plotly_to_base64(fig)
    
    def _plotly_to_base64(self, fig) -> str:
        """Convert Plotly figure to base64 encoded string."""
        try:
            img_bytes = pio.to_image(fig, format="png", engine="kaleido")
            img_base64 = base64.b64encode(img_bytes).decode('utf-8')
            return f"data:image/png;base64,{img_base64}"
        except Exception as e:
            self.logger.warning(f"Error converting plot to base64: {e}")
            return ""
    
    def _generate_html_report(self, report_data: Dict, output_path: Path):
        """Generate HTML report."""
        try:
            template = self._get_html_template()
            html_content = template.render(
                report=report_data,
                config=self.config,
                generation_date=datetime.now().strftime(self.config.date_format)
            )
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.logger.info(f"HTML report generated: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Error generating HTML report: {e}")
            raise
    
    def _generate_pdf_report(self, report_data: Dict, output_path: Path):
        """Generate PDF report using available libraries."""
        try:
            if WEASYPRINT_AVAILABLE:
                self._generate_pdf_with_weasyprint(report_data, output_path)
            elif REPORTLAB_AVAILABLE:
                self._generate_pdf_with_reportlab(report_data, output_path)
            else:
                raise ImportError("No PDF generation library available")
                
        except Exception as e:
            self.logger.error(f"Error generating PDF report: {e}")
            raise
    
    def _generate_pdf_with_weasyprint(self, report_data: Dict, output_path: Path):
        """Generate PDF using WeasyPrint."""
        # Generate HTML first
        html_template = self._get_html_template()
        html_content = html_template.render(
            report=report_data,
            config=self.config,
            generation_date=datetime.now().strftime(self.config.date_format),
            for_pdf=True
        )
        
        # Convert to PDF
        weasyprint.HTML(string=html_content).write_pdf(str(output_path))
        self.logger.info(f"PDF report generated with WeasyPrint: {output_path}")
    
    def _generate_pdf_with_reportlab(self, report_data: Dict, output_path: Path):
        """Generate PDF using ReportLab."""
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=1,  # Center
            spaceAfter=30
        )
        story.append(Paragraph(report_data['metadata']['title'], title_style))
        story.append(Spacer(1, 12))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        exec_summary = report_data.get('executive_summary', {})
        story.append(Paragraph(exec_summary.get('ai_generated_summary', ''), styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Key Metrics Table
        if 'key_metrics' in exec_summary:
            metrics = exec_summary['key_metrics']
            table_data = [['Metric', 'Value']]
            for key, value in metrics.items():
                table_data.append([key.replace('_', ' ').title(), str(value)])
            
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
        
        # Build PDF
        doc.build(story)
        self.logger.info(f"PDF report generated with ReportLab: {output_path}")
    
    def _generate_word_report(self, report_data: Dict, output_path: Path):
        """Generate Word document report."""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(report_data['metadata']['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            exec_summary = report_data.get('executive_summary', {})
            doc.add_paragraph(exec_summary.get('ai_generated_summary', ''))
            
            # Key Metrics
            if 'key_metrics' in exec_summary:
                doc.add_heading('Key Metrics', level=2)
                metrics = exec_summary['key_metrics']
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Metric'
                hdr_cells[1].text = 'Value'
                
                for key, value in metrics.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = key.replace('_', ' ').title()
                    row_cells[1].text = str(value)
            
            # Data Overview
            doc.add_heading('Data Overview', level=1)
            data_overview = report_data.get('data_overview', {})
            basic_info = data_overview.get('basic_info', {})
            
            overview_text = f"""
            The dataset contains {basic_info.get('rows', 'N/A')} rows and {basic_info.get('columns', 'N/A')} columns.
            Memory usage is approximately {basic_info.get('memory_usage', 'N/A')}.
            """
            doc.add_paragraph(overview_text)
            
            # Key Findings
            if 'key_findings' in report_data:
                doc.add_heading('Key Findings', level=1)
                for finding in report_data['key_findings']:
                    doc.add_paragraph(f"• {finding}", style='List Bullet')
            
            # Recommendations
            if 'recommendations' in report_data:
                doc.add_heading('Recommendations', level=1)
                for recommendation in report_data['recommendations']:
                    doc.add_paragraph(f"• {recommendation}", style='List Bullet')
            
            # Save document
            doc.save(str(output_path))
            self.logger.info(f"Word report generated: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Error generating Word report: {e}")
            raise
    
    def _get_html_template(self) -> Template:
        """Get or create HTML template."""
        try:
            return self.template_env.get_template('report_template.html')
        except:
            return Template(self._get_default_html_template())
    
    def _create_default_templates(self):
        """Create default report templates."""
        template_path = Path(self.config.template_dir)
        
        # HTML template
        html_template_path = template_path / 'report_template.html'
        if not html_template_path.exists():
            with open(html_template_path, 'w', encoding='utf-8') as f:
                f.write(self._get_default_html_template())
    
    def _get_default_html_template(self) -> str:
        """Get default HTML template."""
        return """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{{ report.metadata.title }}</title>
            <style>
                body {
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 1200px;
                    margin: 0 auto;
                    padding: 20px;
                    background-color: #f8f9fa;
                }
                .header {
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    padding: 2rem;
                    border-radius: 10px;
                    text-align: center;
                    margin-bottom: 2rem;
                }
                .section {
                    background: white;
                    padding: 1.5rem;
                    margin-bottom: 1.5rem;
                    border-radius: 8px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                }
                .metric-grid {
                    display: grid;
                    grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                    gap: 1rem;
                    margin: 1rem 0;
                }
                .metric-card {
                    background: #f8f9fa;
                    padding: 1rem;
                    border-radius: 6px;
                    text-align: center;
                    border-left: 4px solid #667eea;
                }
                .metric-value {
                    font-size: 1.5rem;
                    font-weight: bold;
                    color: #667eea;
                }
                .visualization {
                    text-align: center;
                    margin: 1rem 0;
                }
                .visualization img {
                    max-width: 100%;
                    height: auto;
                    border-radius: 6px;
                    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
                }
                .insight-list {
                    list-style: none;
                    padding: 0;
                }
                .insight-list li {
                    background: #e3f2fd;
                    margin: 0.5rem 0;
                    padding: 0.75rem;
                    border-radius: 4px;
                    border-left: 3px solid #2196f3;
                }
                .footer {
                    text-align: center;
                    color: #666;
                    font-size: 0.9rem;
                    margin-top: 2rem;
                    padding: 1rem;
                    border-top: 1px solid #ddd;
                }
                h1 { color: #333; margin-bottom: 0.5rem; }
                h2 { color: #667eea; border-bottom: 2px solid #667eea; padding-bottom: 0.5rem; }
                h3 { color: #555; }
                .quality-score {
                    font-size: 2rem;
                    font-weight: bold;
                    text-align: center;
                    padding: 1rem;
                    border-radius: 50%;
                    width: 80px;
                    height: 80px;
                    margin: 0 auto;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                }
                .score-excellent { background: #4caf50; color: white; }
                .score-good { background: #ff9800; color: white; }
                .score-poor { background: #f44336; color: white; }
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{{ report.metadata.title }}</h1>
                <p>Generated by {{ config.company_name }} on {{ generation_date }}</p>
            </div>

            {% if report.executive_summary %}
            <div class="section">
                <h2>Executive Summary</h2>
                <p>{{ report.executive_summary.ai_generated_summary }}</p>
                
                <div class="metric-grid">
                    {% for key, value in report.executive_summary.key_metrics.items() %}
                    <div class="metric-card">
                        <div class="metric-value">{{ value }}</div>
                        <div>{{ key.replace('_', ' ').title() }}</div>
                    </div>
                    {% endfor %}
                </div>
            </div>
            {% endif %}

            {% if report.data_overview %}
            <div class="section">
                <h2>Data Overview</h2>
                <div class="metric-grid">
                    {% for key, value in report.data_overview.basic_info.items() %}
                    <div class="metric-card">
                        <div class="metric-value">{{ value }}</div>
                        <div>{{ key.replace('_', ' ').title() }}</div>
                    </div>
                    {% endfor %}
                </div>
            </div>
            {% endif %}

            {% if report.quality_assessment %}
            <div class="section">
                <h2>Data Quality Assessment</h2>
                <div style="text-align: center; margin: 2rem 0;">
                    {% set score = report.quality_assessment.overall_score %}
                    <div class="quality-score {% if score >= 80 %}score-excellent{% elif score >= 60 %}score-good{% else %}score-poor{% endif %}">
                        {{ "%.0f"|format(score) }}%
                    </div>
                    <p>Overall Quality Score</p>
                </div>
                
                {% if report.quality_assessment.quality_issues %}
                <h3>Quality Issues Identified:</h3>
                <ul class="insight-list">
                    {% for issue in report.quality_assessment.quality_issues %}
                    <li>{{ issue }}</li>
                    {% endfor %}
                </ul>
                {% endif %}
            </div>
            {% endif %}

            {% if report.visualizations %}
            <div class="section">
                <h2>Data Visualizations</h2>
                {% for viz_name, viz_data in report.visualizations.items() %}
                {% if viz_data %}
                <div class="visualization">
                    <h3>{{ viz_name.replace('_', ' ').title() }}</h3>
                    <img src="{{ viz_data }}" alt="{{ viz_name }}">
                </div>
                {% endif %}
                {% endfor %}
            </div>
            {% endif %}

            {% if report.key_findings %}
            <div class="section">
                <h2>Key Findings</h2>
                <ul class="insight-list">
                    {% for finding in report.key_findings %}
                    <li>{{ finding }}</li>
                    {% endfor %}
                </ul>
            </div>
            {% endif %}

            {% if report.recommendations %}
            <div class="section">
                <h2>Recommendations</h2>
                <ul class="insight-list">
                    {% for recommendation in report.recommendations %}
                    <li>{{ recommendation }}</li>
                    {% endfor %}
                </ul>
            </div>
            {% endif %}

            {% if report.model_results %}
            <div class="section">
                <h2>Machine Learning Results</h2>
                <div class="metric-grid">
                    <div class="metric-card">
                        <div class="metric-value">{{ report.model_results.best_model.model_name }}</div>
                        <div>Best Model</div>
                    </div>
                    {% for metric, value in report.model_results.best_model.performance.items() %}
                    <div class="metric-card">
                        <div class="metric-value">{{ "%.3f"|format(value) }}</div>
                        <div>{{ metric.replace('_', ' ').title() }}</div>
                    </div>
                    {% endfor %}
                </div>
            </div>
            {% endif %}

            <div class="footer">
                <p>Report generated automatically by {{ config.company_name }}</p>
                <p>This analysis was created using advanced AI and machine learning techniques</p>
            </div>
        </body>
        </html>
        """
    
    # Helper methods
    def _calculate_quality_score(self, profiling_results: Dict) -> float:
        """Calculate overall data quality score."""
        overview = profiling_results.get('overview', {})
        
        missing_pct = overview.get('missing_data', {}).get('missing_percentage', 0)
        duplicate_pct = overview.get('duplicates', {}).get('percentage', 0)
        
        # Simple scoring algorithm
        completeness_score = max(0, 100 - missing_pct * 2)
        uniqueness_score = max(0, 100 - duplicate_pct * 5)
        
        overall_score = (completeness_score + uniqueness_score) / 2
        return min(100, max(0, overall_score))
    
    def _calculate_completeness_metrics(self, profiling_results: Dict) -> Dict[str, Any]:
        """Calculate data completeness metrics."""
        overview = profiling_results.get('overview', {})
        missing_info = overview.get('missing_data', {})
        
        return {
            'overall_completeness': 100 - missing_info.get('missing_percentage', 0),
            'columns_with_missing': len(missing_info.get('columns_with_missing', [])),
            'complete_records': missing_info.get('complete_records', 0),
            'completeness_grade': self._get_completeness_grade(100 - missing_info.get('missing_percentage', 0))
        }
    
    def _get_completeness_grade(self, completeness_pct: float) -> str:
        """Get grade based on completeness percentage."""
        if completeness_pct >= 95:
            return "Excellent"
        elif completeness_pct >= 85:
            return "Good"
        elif completeness_pct >= 70:
            return "Fair"
        else:
            return "Poor"
    
    def _create_summary_table(self, numeric_summaries: Dict) -> Dict[str, Any]:
        """Create summary statistics table."""
        if not numeric_summaries:
            return {}
        
        summary_table = {
            'columns': list(numeric_summaries.keys()),
            'mean_values': [stats['descriptive'].get('mean', 0) for stats in numeric_summaries.values()],
            'std_values': [stats['descriptive'].get('std', 0) for stats in numeric_summaries.values()],
            'min_values': [stats['descriptive'].get('min', 0) for stats in numeric_summaries.values()],
            'max_values': [stats['descriptive'].get('max', 0) for stats in numeric_summaries.values()]
        }
        
        return summary_table
    
    def _generate_quality_recommendations(self, quality_issues: List[str]) -> List[str]:
        """Generate recommendations based on quality issues."""
        recommendations = []
        
        if not quality_issues:
            recommendations.append("Data quality is excellent - no major issues detected.")
            return recommendations
        
        if any("missing" in issue.lower() for issue in quality_issues):
            recommendations.append("Implement missing value imputation strategies for columns with high missing rates.")
        
        if any("duplicate" in issue.lower() for issue in quality_issues):
            recommendations.append("Remove duplicate records to improve data quality and model performance.")
        
        recommendations.append("Consider data validation rules to prevent future quality issues.")
        recommendations.append("Implement automated data quality monitoring.")
        
        return recommendations
    
    def _get_highest_correlation(self, correlations: Dict) -> float:
        """Get the highest correlation value."""
        strong_correlations = correlations.get('strong_correlations', [])
        if strong_correlations:
            return max(abs(corr.get('pearson_correlation', 0)) for corr in strong_correlations)
        return 0.0
    
    def _assess_multicollinearity_risk(self, correlations: Dict) -> str:
        """Assess multicollinearity risk."""
        strong_correlations = correlations.get('strong_correlations', [])
        high_corr_count = sum(1 for corr in strong_correlations if abs(corr.get('pearson_correlation', 0)) > 0.8)
        
        if high_corr_count > 3:
            return "High"
        elif high_corr_count > 1:
            return "Moderate"
        else:
            return "Low"
    
    def _extract_key_findings(self, ai_insights: Dict) -> List[str]:
        """Extract key findings from AI insights."""
        findings = ai_insights.get('key_findings', [])
        if not findings:
            # Fallback to other insight categories
            all_insights = []
            for category in ['data_quality_insights', 'statistical_insights', 'correlation_insights']:
                category_insights = ai_insights.get(category, [])
                if isinstance(category_insights, list):
                    all_insights.extend(category_insights[:2])  # Top 2 from each category
            findings = all_insights[:5]  # Top 5 overall
        
        return findings
    
    def _extract_recommendations(self, ai_insights: Dict) -> List[str]:
        """Extract recommendations from AI insights."""
        recommendations = ai_insights.get('recommendations', [])
        if not recommendations:
            recommendations = ai_insights.get('actionable_items', [])
        
        return recommendations
    
    def _create_appendix(self, dataset_info: Dict, profiling_results: Dict) -> Dict[str, Any]:
        """Create appendix with technical details."""
        return {
            'technical_details': {
                'analysis_timestamp': datetime.now().isoformat(),
                'platform_version': '1.0.0',
                'analysis_engine': 'NISR Rwanda Analytics Platform'
            },
            'data_dictionary': self._create_data_dictionary(profiling_results),
            'methodology': [
                "Automated data profiling using statistical analysis",
                "AI-powered insight generation using natural language processing",
                "Machine learning model evaluation using cross-validation",
                "Visualization generation using Plotly and Matplotlib"
            ]
        }
    
    def _create_data_dictionary(self, profiling_results: Dict) -> Dict[str, Any]:
        """Create data dictionary from profiling results."""
        columns = profiling_results.get('columns', {})
        data_dictionary = {}
        
        for col_name, col_data in columns.items():
            basic_stats = col_data.get('basic_stats', {})
            data_dictionary[col_name] = {
                'missing_percentage': basic_stats.get('missing_percentage', 0),
                'unique_count': basic_stats.get('unique_count', 0),
                'data_type': 'numeric' if 'numeric_stats' in col_data else 'categorical'
            }
        
        return data_dictionary
    
    def _generate_report_metadata(self, report_data: Dict, generated_reports: Dict) -> Dict[str, Any]:
        """Generate metadata for the report."""
        return {
            'report_info': {
                'title': report_data['metadata']['title'],
                'generated_at': datetime.now().isoformat(),
                'generated_by': self.config.company_name,
                'report_version': '1.0.0'
            },
            'dataset_summary': {
                'shape': report_data['dataset_info'].get('shape', (0, 0)),
                'quality_score': report_data.get('quality_assessment', {}).get('overall_score', 0),
                'analysis_type': 'comprehensive_data_analysis'
            },
            'generated_files': generated_reports,
            'configuration': {
                'include_visualizations': self.config.include_visualizations,
                'include_ai_insights': self.config.include_insights,
                'include_recommendations': self.config.include_recommendations
            }
        }
